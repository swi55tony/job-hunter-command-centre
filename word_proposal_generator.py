# word_proposal_generator.py - Word Document Generator
"""
Professional Word document generator for proposals
Creates properly formatted documents with consistent styling
"""

import os
from datetime import datetime
import re

class WordProposalGenerator:
    def __init__(self):
        # Set up document folder - will use user's home directory or current working directory
        self.proposals_folder = self._get_proposals_folder()
        print(f"✅ Word Generator initialized - saving to: {self.proposals_folder}")
    
    def _get_proposals_folder(self):
        """Get appropriate proposals folder based on environment"""
        try:
            # Try to create proposals folder in user's home directory
            home_dir = os.path.expanduser("~")
            proposals_folder = os.path.join(home_dir, "JobHunter_Proposals")
            os.makedirs(proposals_folder, exist_ok=True)
            return proposals_folder
        except:
            # Fallback to current working directory
            proposals_folder = os.path.join(os.getcwd(), "proposals")
            os.makedirs(proposals_folder, exist_ok=True)
            return proposals_folder
    
    def create_proposal_document(self, job_title, job_url, proposal_text, budget, campaign, score):
        """Create a professional Word document for the proposal"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Document styling
            self._setup_document_style(doc)
            
            # Header section
            self._add_header(doc, job_title, budget, campaign, score)
            
            # Proposal content
            self._add_proposal_content(doc, proposal_text)
            
            # Footer
            self._add_footer(doc, job_url)
            
            # Save document
            filepath = self._save_document(doc, job_title)
            
            if filepath:
                print(f"📄 Word proposal created: {os.path.basename(filepath)}")
                return filepath
            else:
                print("❌ Failed to save Word document")
                return None
                
        except ImportError:
            print("❌ python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            print(f"❌ Word document creation failed: {e}")
            return None
    
    def _setup_document_style(self, doc):
        """Set up document-wide styling"""
        try:
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'  # Use widely available font instead of Poppins
            font.size = Pt(11)
        except:
            pass  # Skip styling if not available
    
    def _add_header(self, doc, job_title, budget, campaign, score):
        """Add professional header to document"""
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('🎖️ ANTONY DRAPER')
        title_run.bold = True
        title_para.alignment = 1  # Center alignment
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run('Executive Operations Consultant | Ex-Military Leadership')
        subtitle_para.alignment = 1
        
        # Separator
        doc.add_paragraph('═' * 80)
        
        # Job details
        details_para = doc.add_paragraph()
        details_para.add_run('OPPORTUNITY: ').bold = True
        details_para.add_run(f'{job_title}\n')
        details_para.add_run('BUDGET: ').bold = True
        details_para.add_run(f'{budget}\n')
        details_para.add_run('CAMPAIGN: ').bold = True
        details_para.add_run(f'{campaign.upper()}\n')
        details_para.add_run('SCORE: ').bold = True
        details_para.add_run(f'{score:.1f}/10\n')
        
        doc.add_paragraph('═' * 80)
    
    def _add_proposal_content(self, doc, proposal_text):
        """Add formatted proposal content"""
        # Section header
        proposal_header = doc.add_paragraph()
        proposal_header.add_run('📋 EXECUTIVE PROPOSAL').bold = True
        
        # Process proposal text
        paragraphs = proposal_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                if '•' in para_text:
                    # Handle bullet points
                    lines = para_text.split('\n')
                    for line in lines:
                        if line.strip():
                            if line.strip().startswith('•'):
                                # Create proper Word bullet point
                                clean_text = line.strip()[1:].strip()
                                try:
                                    bullet_para = doc.add_paragraph(clean_text, style='List Bullet')
                                except:
                                    # Fallback if List Bullet style not available
                                    bullet_para = doc.add_paragraph(f'• {clean_text}')
                            else:
                                # Section header
                                header_para = doc.add_paragraph(line.strip())
                                header_para.runs[0].bold = True
                else:
                    # Regular paragraph
                    doc.add_paragraph(para_text.strip())
    
    def _add_footer(self, doc, job_url):
        """Add footer with job details"""
        doc.add_paragraph('═' * 80)
        footer_para = doc.add_paragraph()
        footer_para.add_run(f'Job URL: {job_url}\n')
        footer_para.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    def _save_document(self, doc, job_title):
        """Save document with clean filename"""
        try:
            # Generate clean filename
            clean_title = re.sub(r'[^\w\s-]', '', job_title)
            clean_title = re.sub(r'\s+', '_', clean_title)[:40]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Proposal_{clean_title}_{timestamp}.docx"
            filepath = os.path.join(self.proposals_folder, filename)
            
            # Save document
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"❌ Error saving document: {e}")
            return None
    
    def create_summary_document(self, proposals_list):
        """Create summary document of all proposals"""
        try:
            from docx import Document
            
            doc = Document()
            
            # Title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run('🎖️ PROPOSAL SUMMARY - ANTONY DRAPER')
            title_run.bold = True
            title_para.alignment = 1
            
            # Stats
            doc.add_paragraph('═' * 80)
            stats_para = doc.add_paragraph()
            stats_para.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
            stats_para.add_run(f'Total Proposals: {len(proposals_list)}\n')
            
            # High-value proposals
            high_value = [p for p in proposals_list if p.get('score', 0) >= 7.0]
            if high_value:
                stats_para.add_run(f'High-Value Proposals (7.0+): {len(high_value)}\n')
            
            doc.add_paragraph('═' * 80)
            
            # List all proposals
            header_para = doc.add_paragraph()
            header_para.add_run('📋 PROPOSAL LIST').bold = True
            
            for i, proposal in enumerate(proposals_list, 1):
                entry_para = doc.add_paragraph()
                
                # Job title (bold)
                title_run = entry_para.add_run(f'{i}. {proposal.get("job_title", "Unknown Job")}\n')
                title_run.bold = True
                
                # Details
                entry_para.add_run(f'   📄 File: {proposal.get("filename", "Unknown")}\n')
                entry_para.add_run(f'   🎯 Score: {proposal.get("score", 0):.1f}/10\n')
                entry_para.add_run(f'   💰 Budget: {proposal.get("budget", "Not specified")}\n')
                entry_para.add_run(f'   📊 Campaign: {proposal.get("campaign", "Unknown")}\n')
                
                # Add space between entries
                doc.add_paragraph()
            
            # Save summary
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Proposal_Summary_{timestamp}.docx"
            filepath = os.path.join(self.proposals_folder, filename)
            
            doc.save(filepath)
            print(f"📄 Proposal summary created: {filename}")
            return filepath
            
        except ImportError:
            print("❌ python-docx not installed for summary creation")
            return None
        except Exception as e:
            print(f"❌ Summary creation failed: {e}")
            return None
    
    def get_proposals_folder(self):
        """Get the current proposals folder path"""
        return self.proposals_folder
    
    def set_proposals_folder(self, folder_path):
        """Set a custom proposals folder"""
        try:
            os.makedirs(folder_path, exist_ok=True)
            self.proposals_folder = folder_path
            print(f"✅ Proposals folder updated: {folder_path}")
            return True
        except Exception as e:
            print(f"❌ Failed to set proposals folder: {e}")
            return False

# Export the class
__all__ = ['WordProposalGenerator']
